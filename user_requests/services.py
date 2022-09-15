from .serializers import UserRequestSerializer
from rest_framework import viewsets, permissions, pagination
from .models import UserRequest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
import pyminizip
import smtplib
from email import encoders
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.utils import formatdate
import configparser


class StandardResultsSetPagination(pagination.PageNumberPagination):
    page_size = 100
    page_size_query_param = 'page_size'
    max_page_size = 10000

class UserRequestViewSet(viewsets.ModelViewSet):
    queryset = UserRequest.objects.order_by('-create_date', '-change_date')
    serializer_class = UserRequestSerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = StandardResultsSetPagination

def get_dictionaries():
    data = {
        'statuses': UserRequest.STATUS,
        'managements': UserRequest.MANAGEMENT,
        'departments': UserRequest.DEPARTMENT,
        'dbs': UserRequest.DB
    }
    return data

def check_iin(id, status, db, iin):
    cnt = UserRequest.objects.filter(status=status, db=db, iin=iin).exclude(id=id).count()
    return { 'result': cnt }

def check_phone(id, status, db, phone):
    cnt = UserRequest.objects.filter(status=status,db=db,phone=phone).exclude(id=id).count()
    return { 'result': cnt }


def create_docx(file_name, ids):
    int_ids = [int(id) for id in ids]
    
    userRequests = UserRequest.objects.filter(id__in=int_ids)

    n = 1
    items = []
    for item in userRequests.all():
        new_item = {}
        new_item['#'] = str(n)
        
        department = ''
        for dep in item.DEPARTMENT:
            if dep[0] == item.department:
                department = dep[1]
        new_item['department'] = department

        management = ''
        for man in item.MANAGEMENT:
            if man[0] == item.management:
                management = man[1]
        new_item['management'] = management


        new_item['full_name'] = f'{item.first_name} {item.middle_name} {item.last_name}' 
        
        dbname = ''
        for db in item.DB:
            if db[0] == item.db:
                dbname = db[1]

        new_item['dbname'] = dbname
        new_item['login'] = str(item.login)
        new_item['password'] = str(item.password)
        items.append(new_item)
        n += 1

    doc = Document()
    section = doc.sections[0]
    section.page_height = Mm(210)
    section.page_width = Mm(297)
    section.left_margin = Mm(30.0)
    section.right_margin = Mm(10.0)
    section.top_margin = Mm(20.0)
    section.bottom_margin = Mm(20.0)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    header = doc.add_heading('Пароли пользователей')
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    
    table = doc.add_table(rows=1, cols=7)
    heading_cells = table.rows[0].cells
    heading_cells[0].text = '№'   
    heading_cells[1].text = 'Подразделение'   
    heading_cells[2].text = 'Служба'   
    heading_cells[3].text = 'Ф.И.О.'
    heading_cells[4].text = 'База данных'
    heading_cells[5].text = 'Логин'
    heading_cells[6].text = 'Пароль'

    for item in items:
        cells = table.add_row().cells
        cells[0].text = item['#']
        cells[1].text = item['department']
        cells[2].text = item['management']
        cells[3].text = item['full_name']
        cells[4].text = item['dbname']
        cells[5].text = item['login']
        cells[6].text = item['password']

    table.style = 'Table Grid'

    doc.add_page_break()

    header = doc.add_heading('Ведомость на получение паролей')
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=7)
    heading_cells = table.rows[0].cells
    heading_cells[0].text = '№'   
    heading_cells[1].text = 'Подразделение'   
    heading_cells[2].text = 'Служба'  
    heading_cells[3].text = 'Ф.И.О.'
    heading_cells[4].text = 'База данных'
    heading_cells[5].text = 'Дата получения'
    heading_cells[6].text = 'Подпись'

    for item in items:
        cells = table.add_row().cells
        cells[0].text = item['#']
        cells[1].text = item['department']
        cells[2].text = item['management']
        cells[3].text = item['full_name']
        cells[4].text = item['dbname']
        
    table.style = 'Table Grid'
    doc.save(file_name)

def send_email(server, port, login, password, from_addr, to_addr, subject, file_to_attach):
    msg = MIMEMultipart()
    msg["From"] = from_addr
    msg["Subject"] = subject
    msg["Date"] = formatdate(localtime=True)
    
    attachment = MIMEBase('application', "octet-stream")
    header = 'Content-Disposition', f'attachment; filename="{file_to_attach}"'
    try:
        with open(file_to_attach, "rb") as fh:
            data = fh.read()
        attachment.set_payload(data)
        encoders.encode_base64(attachment)
        attachment.add_header(*header)
        msg.attach(attachment)
    except IOError:
        print(f"Ошибка при открытии файла вложения {file_to_attach}")
    
    try:
        smtp = smtplib.SMTP(server, port)
        # smtp.starttls()
        # smtp.ehlo()
        smtp.login(login, password)
        smtp.sendmail(from_addr, to_addr, msg.as_string())
    except smtplib.SMTPException as err:
        print('Что - то пошло не так...')
        raise err
    finally:
        smtp.quit()


def get_report(ids):
    config = configparser.ConfigParser()
    config.sections()
    config.read('config/config.ini')

    doc_file = config['DOCX']['FileName']
    
    create_docx(doc_file, ids)

    file = open(doc_file, 'rb')
    
    return file

def send_report(ids, to_addr, zip_pas):
    config = configparser.ConfigParser()
    config.read('config/config.ini', encoding='utf-8')

    doc_file = config['DOCX']['FileName']
    zip_file = config['ZIP']['FileName']
    compress_level = config['ZIP']['CompressLevel']
    server = config['EMAIL']['Server']
    port = config['EMAIL']['Port']
    login = config['EMAIL']['Login']
    password = config['EMAIL']['Password']
    from_addr = config['EMAIL']['Addr']
    subject = config['EMAIL']['Subject']

    create_docx(doc_file, ids)

    pyminizip.compress(doc_file, '/', zip_file, zip_pas, int(compress_level))

    send_email(server, port, login, password, from_addr, to_addr, subject, zip_file)

def get_addr_list():
    config = configparser.ConfigParser()
    config.read('config/config.ini', encoding='utf-8')
    
    addr_list = []
    for key in config['ADDR_LIST']:
        addr_list.append((config['ADDR_LIST'][key], key))
    
    data = {'addrList': addr_list}   
    return data